"""Generate Norrsken membership agreement docx from markdown."""
from __future__ import annotations

import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
MD = os.path.join(HERE, "Norrsken-Membership-Agreement-CON-344.md")
OUT = os.path.join(HERE, "Norrsken-Membership-Agreement-CON-344.docx")


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, rows: list[list[str]], header: bool = True) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            cell = table.rows[i].cells[j]
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", text.strip())
            cell.text = clean
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
            if header and i == 0:
                set_cell_shading(cell, "E8EEF4")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
    doc.add_paragraph()


def add_rich_paragraph(doc: Document, text: str, *, quote: bool = False, bullet: bool = False) -> None:
    if bullet:
        para = doc.add_paragraph(style="List Bullet")
    else:
        para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    if quote:
        para.paragraph_format.left_indent = Inches(0.35)

    parts = re.split(r"(\*\*.+?\*\*|\*(.+?)\*)", text)
    for part in parts:
        if not part or part.startswith("("):
            continue
        m_bold = re.match(r"\*\*(.+?)\*\*", part)
        m_italic = re.match(r"\*(.+?)\*", part) if not m_bold else None
        if m_bold:
            run = para.add_run(m_bold.group(1))
            run.bold = True
        elif m_italic:
            run = para.add_run(m_italic.group(1))
            run.italic = True
        else:
            run = para.add_run(part)
        run.font.size = Pt(11)
        if quote:
            run.italic = True


def parse_md_table(lines: list[str]) -> list[list[str]] | None:
    if len(lines) < 2 or "|" not in lines[0]:
        return None
    if not re.match(r"^\|?\s*:?-+", lines[1]):
        return None
    rows: list[list[str]] = []
    for line in lines:
        if "|" not in line:
            break
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def build() -> None:
    with open(MD, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1.1)
        s.right_margin = Inches(1.1)

    i = 0
    while i < len(lines):
        line = lines[i]

        if line.strip() == "---":
            i += 1
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[2:].strip())
            r.bold = True
            r.font.size = Pt(16)
            i += 1
            continue

        if line.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(line[3:].strip())
            r.bold = True
            r.font.size = Pt(13)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        if line.startswith("|"):
            block: list[str] = []
            while i < len(lines) and lines[i].startswith("|"):
                block.append(lines[i])
                i += 1
            rows = parse_md_table(block)
            if rows:
                add_table(doc, rows)
            continue

        if line.startswith("- "):
            add_rich_paragraph(doc, line[2:].strip(), bullet=True)
            i += 1
            continue

        if line.strip().startswith("*") and line.strip().endswith("*") and not line.strip().startswith("**"):
            p = doc.add_paragraph(line.strip().strip("*"))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.italic = True
                r.font.size = Pt(10)
            i += 1
            continue

        if line.strip():
            add_rich_paragraph(doc, line.strip())
        i += 1

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
