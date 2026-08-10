"""Generate JMT-S / Salanor partner agreement docx."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT = os.path.join(os.path.dirname(__file__), "JMTS-Salanor-Contrat-Partenariat-Aegis.docx")

doc = Document()
for s in doc.sections:
    s.top_margin = Inches(1)
    s.bottom_margin = Inches(1)
    s.left_margin = Inches(1.1)
    s.right_margin = Inches(1.1)


def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)


def h(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def p(text):
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    for r in para.runs:
        r.font.size = Pt(11)
    return para


def li(text):
    para = doc.add_paragraph(text, style="List Bullet")
    para.paragraph_format.space_after = Pt(2)
    for r in para.runs:
        r.font.size = Pt(11)


title("CONTRAT DE PARTENARIAT COMMERCIAL")
p_center = doc.add_paragraph()
p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_center.add_run("Produit Aegis — Salanor Ltd / JMT-S")
r.font.size = Pt(12)
doc.add_paragraph()
p("Version 1.0 — [Date]")
doc.add_paragraph()

p("ENTRE LES SOUSSIGNÉS :")
doc.add_paragraph()
p("SALANOR Ltd")
p("Société de droit rwandais, immatriculée au Rwanda Development Board,")
p("TIN : 152357888,")
p("Siège : Kabeza Hirwa, Muhima, Nyarugenge, Kigali, Rwanda,")
p("Représentée par Joseph Landry Bougang Fotso, Managing Director,")
p("Email : landry@salanor.com")
p("Ci-après « Salanor » ou « le Fournisseur ».")
doc.add_paragraph()
p("ET :")
doc.add_paragraph()
p("[DÉNOMINATION SOCIALE JMT-S]")
p("[Forme juridique], RCS [Ville] n° [●], SIRET : [●],")
p("Siège : [Adresse], France,")
p("Représentée par Jacob [Nom de famille], [Qualité],")
p("Email : [●]")
p("Ci-après « le Partenaire ».")
doc.add_paragraph()
p("Ci-après ensemble « les Parties ».")

h("PRÉAMBULE")
p(
    "Salanor édite Aegis, plateforme d’audit et de traçabilité des actions des agents "
    "d’intelligence artificielle en environnement professionnel."
)
p(
    "Le Partenaire souhaite promouvoir et commercialiser Aegis en France. "
    "Les Parties conviennent des conditions de leur collaboration commerciale "
    "dans le cadre du présent contrat."
)

h("ARTICLE 1 — DÉFINITIONS")
p("Aux fins du contrat :")
li("« Aegis » : la solution logicielle de Salanor (console, API, SDK, hébergement et services associés).")
li("« Client attribué » : client dont l’opportunité a été enregistrée conformément à l’article 4 et qui a souscrit un abonnement payant à Aegis.")
li("« Revenu net d’abonnement » : montants encaissés par Salanor au titre de l’abonnement Aegis d’un Client attribué, hors taxes, remboursements et impayés.")
li("« Opportunité » : prospect identifié et présenté par le Partenaire avant conclusion d’un contrat Aegis avec ce prospect.")

h("ARTICLE 2 — OBJET")
p(
    "Salanor confère au Partenaire un droit non exclusif de promotion et de revente "
    "des abonnements Aegis sur le Territoire, moyennant commission selon l’article 5."
)
p(
    "Le Partenaire agit en qualité de partenaire commercial indépendant. Le présent "
    "contrat ne crée ni société commune, ni mandat, ni relation de subordination, "
    "ni contrat de travail."
)

h("ARTICLE 3 — TERRITOIRE")
p("3.1 Le Territoire couvre la République française (métropole et DOM le cas échéant).")
p("3.2 La collaboration est non exclusive. Salanor peut vendre directement ou via d’autres partenaires en France.")
p("3.3 Toute extension à d’autres pays de l’Union européenne requiert un avenant écrit.")

h("ARTICLE 4 — ENREGISTREMENT DES OPPORTUNITÉS")
p(
    "4.1 Toute Opportunité est enregistrée par écrit à partners@salanor.com avec : "
    "raison sociale, contact, date de prise de contact, description du besoin."
)
p("4.2 Salanor confirme sous cinq (5) jours ouvrés si l’Opportunité est libre ou déjà prise.")
p("4.3 Protection de quatre-vingt-dix (90) jours à compter de la confirmation, renouvelable une fois sur demande motivée.")
p("4.4 Aucune commission n’est due si :")
li("le prospect était déjà en relation avec Salanor avant enregistrement ;")
li("le client est venu en direct (site, email, événement) sans intervention préalable du Partenaire ;")
li("un autre partenaire a enregistré l’Opportunité en premier ;")
li("le client est hors Territoire, sauf accord écrit pour ce dossier.")
p("4.5 Le registre tenu par Salanor fait foi, sous réserve des preuves du Partenaire.")

h("ARTICLE 5 — COMMISSION")
p("5.1 Taux : vingt pour cent (20 %) du Revenu net d’abonnement pour chaque Client attribué.")
p(
    "5.2 Durée : la commission est due pendant vingt-quatre (24) mois à compter de la "
    "première facture payée du Client. Passé ce délai, aucune commission sur renouvellements."
)
p("5.3 Upsell : pendant ces 24 mois, la commission s’applique aux extensions d’abonnement du même client.")
p(
    "5.4 Apport d’affaires simple : si le Partenaire se limite à l’introduction sans vente "
    "ni support L1, un taux de quinze pour cent (15 %) peut s’appliquer par accord écrit "
    "pour ce client. À défaut, 20 % s’applique."
)
p("5.5 Paiement trimestriel, sous trente (30) jours après fin de trimestre, sur encaissements effectifs uniquement.")
p("5.6 Aucun revenu minimum garanti.")

h("ARTICLE 6 — PRIX ET CONTRATS CLIENTS")
p("6.1 Salanor fixe les tarifs. Remise maximale du Partenaire : 10 % sans accord écrit.")
p("6.2 Sauf accord contraire, l’abonnement Aegis est conclu entre Salanor et le Client attribué.")
p("6.3 Le Partenaire ne peut engager Salanor envers un client final sans mandat écrit préalable de Salanor.")

h("ARTICLE 7 — OBLIGATIONS DE SALANOR")
li("Documentation commerciale et accès démo (sandbox) ;")
li("Support produit niveau 2 (L2) ;")
li("Information des changements majeurs affectant le produit ou les commissions.")

h("ARTICLE 8 — OBLIGATIONS DU PARTENAIRE")
li("Promotion loyale conforme à la documentation Salanor ;")
li("Pas d’engagements au-delà des documents approuvés ;")
li("Support commercial et L1 en français pour les clients qu’il vend activement ;")
li("Respect du RGPD et des règles de prospection B2B ;")
li("Enregistrement systématique des Opportunités.")

h("ARTICLE 9 — PROPRIÉTÉ INTELLECTUELLE")
p("9.1 Aegis, le code, les marques et la documentation restent la propriété exclusive de Salanor.")
p("9.2 Aucune cession de droits de propriété intellectuelle au Partenaire.")
p("9.3 Usage des marques Salanor limité à la promotion d’Aegis pendant la durée du contrat.")

h("ARTICLE 10 — DONNÉES PERSONNELLES")
p(
    "Chaque Partie traite les données dont elle est responsable conformément au RGPD. "
    "Lorsque le Partenaire traite des données pour le compte de Salanor ou d’un client "
    "Aegis, les Parties concluront un accord de sous-traitance (DPA) si nécessaire."
)

h("ARTICLE 11 — CONFIDENTIALITÉ")
p(
    "Les informations commerciales, techniques et financières échangées sont confidentielles "
    "pendant le contrat et cinq (5) ans après son terme, sauf obligation légale ou "
    "information publique."
)

h("ARTICLE 12 — DURÉE ET RÉSILIATION")
p("12.1 Durée initiale : douze (12) mois à compter de la signature.")
p("12.2 Renouvellement tacite par périodes de douze (12) mois, sauf dénonciation avec préavis de quatre-vingt-dix (90) jours.")
p("12.3 Résiliation immédiate en cas de manquement grave non corrigé sous trente (30) jours après mise en demeure.")
p(
    "12.4 À la résiliation, les commissions restent dues pour les Clients attribués jusqu’à "
    "expiration de leur période de 24 mois, sous réserve des encaissements effectifs."
)

h("ARTICLE 13 — RESPONSABILITÉ")
p(
    "Chaque Partie est responsable de ses propres actes. La responsabilité cumulée de "
    "Salanor au titre du présent contrat est limitée au montant total des commissions "
    "versées au Partenaire sur les douze (12) mois précédant le fait générateur."
)

h("ARTICLE 14 — DROIT APPLICABLE — LITIGES")
p(
    "Le contrat est régi par le droit français. À défaut d’accord amiable sous soixante (60) jours, "
    "compétence exclusive est attribuée au Tribunal de commerce de [Ville du siège du Partenaire], "
    "sous réserve des règles impératives applicables à Salanor en Rwanda."
)

h("ARTICLE 15 — DISPOSITIONS FINALES")
p("15.1 Avenant écrit pour toute modification.")
p("15.2 Nullité partielle sans affecter le reste du contrat.")
p("15.3 Signatures électroniques admises.")

doc.add_paragraph()
doc.add_paragraph()
p("Fait à [Lieu], le [Date], en deux exemplaires.")
doc.add_paragraph()
p("Pour Salanor Ltd")
doc.add_paragraph()
p("_______________________________")
p("Joseph Landry Bougang Fotso")
p("Managing Director")
doc.add_paragraph()
doc.add_paragraph()
p("Pour JMT-S")
doc.add_paragraph()
p("_______________________________")
p("Jacob [Nom de famille]")
p("[Qualité]")

doc.save(OUT)
print(OUT)
